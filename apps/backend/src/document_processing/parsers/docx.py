"""DOCX document parser using python-docx."""

import logging
import time
from datetime import datetime, timezone

from shared.models.processing import (
    DocumentMetadata,
    Page,
    ParsedDocument,
    ProcessingStatistics,
    ProcessingWarning,
)
from src.document_processing.exceptions import ParseError
from src.document_processing.parsers.base import BaseParser

logger = logging.getLogger(__name__)

try:
    from docx import Document as DocxDocument
except ImportError as exc:
    raise ImportError(
        "python-docx is required for DOCX parsing. "
        "Install it with: pip install python-docx"
    ) from exc


class DOCXParser(BaseParser):
    """Parse DOCX documents and extract text content using python-docx."""

    @property
    def name(self) -> str:
        return "docx"

    @classmethod
    def supported_formats(cls) -> list[str]:
        return ["docx"]

    async def parse(self, file_path: str) -> ParsedDocument:
        """Extract text from a DOCX file.

        Args:
            file_path: Absolute path to the DOCX file.

        Returns:
            A ParsedDocument with extracted text and metadata.

        Raises:
            ParseError: If the document cannot be opened or parsed.
        """
        start_time = time.monotonic()
        warnings: list[ProcessingWarning] = []

        # --- Open the document ---
        try:
            doc = DocxDocument(file_path)
        except Exception as exc:
            raise ParseError(
                message=f"Failed to open DOCX: {exc}",
                details={"file_path": file_path, "error": str(exc)},
            ) from exc

        try:
            # --- Extract paragraphs ---
            paragraphs = [p.text for p in doc.paragraphs]
            raw_text = "\n".join(paragraphs)

            # Edge case: empty document
            if not raw_text.strip():
                warnings.append(
                    ProcessingWarning(
                        stage="parse",
                        message="DOCX document contains no extractable text",
                        details={"file_path": file_path},
                    )
                )

            char_count = len(raw_text)
            word_count = len(raw_text.split()) if raw_text.strip() else 0

            # Build pages: treat the full document body as a single page
            pages = [
                Page(
                    number=1,
                    content=raw_text,
                    char_count=char_count,
                    word_count=word_count,
                )
            ]

            # --- Extract core properties (never fails; missing = None) ---
            cp = doc.core_properties

            creation_date = _utc_from_naive(cp.created)
            modification_date = _utc_from_naive(cp.modified)

            keywords: list[str] = []
            if cp.keywords:
                raw_keywords = cp.keywords
                if isinstance(raw_keywords, str):
                    keywords = [
                        k.strip()
                        for k in raw_keywords.replace(";", ",").split(",")
                        if k.strip()
                    ]

            # Build extra metadata from available core properties
            extra: dict = {"format": "DOCX"}
            if cp.revision is not None:
                extra["revision"] = cp.revision
            if cp.category:
                extra["category"] = cp.category
            if cp.content_status:
                extra["content_status"] = cp.content_status
            if cp.identifier:
                extra["identifier"] = cp.identifier
            if cp.version:
                extra["version"] = cp.version

            metadata = DocumentMetadata(
                title=cp.title or None,
                author=cp.author or None,
                subject=cp.subject or None,
                keywords=keywords,
                page_count=1,
                word_count=word_count,
                char_count=char_count,
                creation_date=creation_date,
                modification_date=modification_date,
                processed_by=self.name,
                extra=extra,
            )

            # --- Build ProcessingStatistics ---
            elapsed = (time.monotonic() - start_time) * 1000
            statistics = ProcessingStatistics(
                parse_time_ms=elapsed,
                total_time_ms=elapsed,
                page_count=1,
                raw_char_count=char_count,
                clean_char_count=char_count,
            )

            return ParsedDocument(
                parser_used=self.name,
                metadata=metadata,
                pages=pages,
                raw_text=raw_text,
                clean_text=raw_text,
                warnings=warnings,
                statistics=statistics,
            )

        except ParseError:
            raise
        except Exception as exc:
            raise ParseError(
                message=f"Failed to parse DOCX content: {exc}",
                details={"file_path": file_path, "error": str(exc)},
            ) from exc


# ---------------------------------------------------------------------------
# Helper utilities
# ---------------------------------------------------------------------------


def _utc_from_naive(dt: datetime | None) -> datetime | None:
    """Return a UTC-aware datetime from a potentially naive one, or None."""
    if dt is None:
        return None
    try:
        return dt if dt.tzinfo is not None else dt.replace(tzinfo=timezone.utc)
    except (ValueError, AttributeError, OverflowError):
        return None
